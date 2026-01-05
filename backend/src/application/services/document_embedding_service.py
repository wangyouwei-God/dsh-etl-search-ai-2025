"""
Application Service: Document Embedding for RAG

This module provides functionality for extracting text from supporting documents
(PDFs, DOCX, TXT) and generating vector embeddings for RAG (Retrieval-Augmented Generation).

Author: University of Manchester RSE Team
"""

import logging
import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from uuid import uuid4

logger = logging.getLogger(__name__)

# Try to import PDF extraction libraries
try:
    from pypdf import PdfReader as PyPdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

# Try to import DOCX extraction library
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. DOCX extraction will be limited.")


@dataclass
class DocumentChunk:
    """A chunk of document content for embedding."""
    id: str
    document_id: str
    dataset_id: str
    chunk_index: int
    content: str
    source_file: str
    document_type: str
    created_at: datetime

    def to_embedding_text(self) -> str:
        """Format content for embedding."""
        return self.content


class DocumentTextExtractor:
    """
    Extracts text content from various document formats.
    
    Supports:
    - PDF (via PyPDF2)
    - TXT (plain text)
    - DOCX (basic text extraction)
    """
    
    # Maximum characters per chunk (roughly 500 tokens)
    DEFAULT_CHUNK_SIZE = 1500
    CHUNK_OVERLAP = 200
    
    def __init__(self, chunk_size: int = None, overlap: int = None):
        self.chunk_size = chunk_size or self.DEFAULT_CHUNK_SIZE
        self.overlap = overlap or self.CHUNK_OVERLAP
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted text or None if extraction fails
        """
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        ext = path.suffix.lower()
        
        try:
            if ext == '.pdf':
                return self._extract_pdf(path)
            elif ext == '.txt':
                return self._extract_txt(path)
            elif ext in ('.doc', '.docx'):
                return self._extract_docx(path)
            else:
                logger.warning(f"Unsupported file format: {ext}")
                return None
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            return None
    
    def _extract_pdf(self, path: Path) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            if HAS_PYPDF:
                reader = PyPdfReader(str(path))
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            elif HAS_PYPDF2:
                text_parts = []
                with open(path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
            else:
                logger.warning("No PDF library available for extraction")
                return None

            full_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} chars from PDF: {path.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return None

    def _extract_txt(self, path: Path) -> Optional[str]:
        """Extract text from plain text file."""
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            logger.info(f"Read {len(text)} chars from TXT: {path.name}")
            return text
        except Exception as e:
            logger.error(f"TXT read error: {e}")
            return None
    
    def _extract_docx(self, path: Path) -> Optional[str]:
        """Extract text from DOCX file."""
        if not HAS_DOCX:
            logger.warning("python-docx not available for DOCX extraction")
            return None
        if path.suffix.lower() != '.docx':
            logger.warning(f"Unsupported DOC format: {path.name}")
            return None

        try:
            doc = DocxDocument(str(path))
            parts = [p.text for p in doc.paragraphs if p.text]

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            parts.append(cell_text)

            text = "\n".join(parts)
            logger.info(f"Extracted {len(text)} chars from DOCX: {path.name}")
            return text
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return None

    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks for embedding.
        
        Args:
            text: Full document text
            
        Returns:
            List of text chunks
        """
        if not text or len(text) < self.chunk_size:
            return [text] if text else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence or paragraph boundary
            if end < len(text):
                # Look for paragraph break
                para_break = text.rfind('\n\n', start, end)
                if para_break > start + self.chunk_size // 2:
                    end = para_break
                else:
                    # Look for sentence break
                    sentence_break = max(
                        text.rfind('. ', start, end),
                        text.rfind('.\n', start, end)
                    )
                    if sentence_break > start + self.chunk_size // 2:
                        end = sentence_break + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.overlap
        
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks


class DocumentEmbeddingService:
    """
    Service for creating embeddings from supporting documents.
    
    This service:
    1. Extracts text from documents (PDF, TXT)
    2. Chunks text for embedding
    3. Generates embeddings using the embedding service
    4. Stores embeddings in the vector database
    """
    
    def __init__(
        self,
        embedding_service,  # IEmbeddingService
        vector_repository,  # IVectorRepository
        chunk_size: int = 1500
    ):
        """
        Initialize document embedding service.
        
        Args:
            embedding_service: Service for generating embeddings
            vector_repository: Repository for storing vectors
            chunk_size: Characters per chunk
        """
        self.embedding_service = embedding_service
        self.vector_repository = vector_repository
        self.extractor = DocumentTextExtractor(chunk_size=chunk_size)
    
    def process_document(
        self,
        file_path: str,
        dataset_id: str,
        document_type: str = "supporting_doc"
    ) -> List[DocumentChunk]:
        """
        Process a document and store its embeddings.
        
        Args:
            file_path: Path to the document
            dataset_id: Associated dataset ID
            document_type: Type of document
            
        Returns:
            List of created DocumentChunks
        """
        # Extract text
        text = self.extractor.extract_text(file_path)
        if not text:
            logger.warning(f"No text extracted from {file_path}")
            return []
        
        # Chunk text
        chunks_text = self.extractor.chunk_text(text)
        if not chunks_text:
            return []
        
        # Generate document ID
        doc_id = str(uuid4())
        filename = Path(file_path).name
        
        chunks = []
        for idx, chunk_text in enumerate(chunks_text):
            chunk = DocumentChunk(
                id=f"{doc_id}_chunk_{idx}",
                document_id=doc_id,
                dataset_id=dataset_id,
                chunk_index=idx,
                content=chunk_text,
                source_file=filename,
                document_type=document_type,
                created_at=datetime.utcnow()
            )
            chunks.append(chunk)
            
            # Generate embedding
            try:
                embedding = self.embedding_service.generate_embedding(chunk_text)
                
                # Store in vector database with metadata
                self.vector_repository.upsert_vector(
                    id=chunk.id,
                    vector=embedding,
                    metadata={
                        "title": f"{filename} - Chunk {idx + 1}",
                        "abstract": chunk_text[:500],
                        "keywords": [document_type, "supporting_document"],
                        "type": "document",
                        "dataset_id": dataset_id,
                        "source_file": filename,
                        "chunk_index": idx
                    }
                )
                
                logger.debug(f"Stored embedding for chunk {idx + 1}/{len(chunks_text)}")
                
            except Exception as e:
                logger.error(f"Failed to embed chunk {idx}: {e}")
        
        logger.info(f"Processed {len(chunks)} chunks from {filename}")
        return chunks
    
    def process_directory(
        self,
        directory: str,
        dataset_id: str,
        extensions: List[str] = None
    ) -> List[DocumentChunk]:
        """
        Process all documents in a directory.
        
        Args:
            directory: Directory containing documents
            dataset_id: Associated dataset ID
            extensions: File extensions to process (default: pdf, txt)
            
        Returns:
            List of all created DocumentChunks
        """
        extensions = extensions or ['.pdf', '.txt']
        dir_path = Path(directory)
        
        if not dir_path.exists():
            logger.error(f"Directory not found: {directory}")
            return []
        
        all_chunks = []
        
        for ext in extensions:
            for file_path in dir_path.glob(f"**/*{ext}"):
                chunks = self.process_document(
                    str(file_path),
                    dataset_id,
                    "supporting_doc"
                )
                all_chunks.extend(chunks)
        
        logger.info(f"Processed {len(all_chunks)} total chunks from {directory}")
        return all_chunks
