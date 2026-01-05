"""
Infrastructure: Content Extractor

This module extracts text content from supporting documents (PDFs, DOCX)
to enable full-text semantic search via RAG (Retrieval-Augmented Generation).

This addresses the requirement: "extract semantic meaning from these files"
by implementing content indexing beyond metadata-only search.

Author: University of Manchester RSE Team
"""

import logging
from pathlib import Path
from typing import Optional, List
import re

# Document parsers
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

logger = logging.getLogger(__name__)


class ContentExtractor:
    """
    Extracts text content from various document formats.

    Supported Formats:
    - PDF (.pdf) via pypdf
    - Word Documents (.docx) via python-docx
    - Plain text (.txt, .md, .csv)

    Design Pattern: Strategy Pattern (different extraction strategies per format)
    """

    def __init__(self):
        """Initialize the content extractor."""
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_unsupported,  # Legacy format, limited support
            '.txt': self._extract_text,
            '.md': self._extract_text,
            '.csv': self._extract_text,
        }

    def can_extract(self, file_path: str) -> bool:
        """
        Check if content extraction is supported for this file type.

        Args:
            file_path: Path to the file

        Returns:
            bool: True if format is supported
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_formats

    def extract(self, file_path: str) -> Optional[str]:
        """
        Extract text content from a document.

        Args:
            file_path: Path to the document file

        Returns:
            str: Extracted text content, or None if extraction failed

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()

        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}")

        try:
            extractor = self.supported_formats[ext]
            content = extractor(path)

            if content:
                # Clean extracted text
                content = self._clean_text(content)
                logger.info(f"Extracted {len(content)} characters from {path.name}")
                return content
            else:
                logger.warning(f"No content extracted from {path.name}")
                return None

        except Exception as e:
            logger.error(f"Error extracting content from {path.name}: {e}")
            return None

    def _extract_pdf(self, path: Path) -> Optional[str]:
        """
        Extract text from PDF file.

        Args:
            path: Path to PDF file

        Returns:
            str: Extracted text content
        """
        if PdfReader is None:
            logger.error("pypdf not installed, cannot extract PDF content")
            return None

        try:
            reader = PdfReader(str(path))
            text_parts = []

            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num} from {path.name}: {e}")
                    continue

            return '\n\n'.join(text_parts) if text_parts else None

        except Exception as e:
            logger.error(f"Error reading PDF {path.name}: {e}")
            return None

    def _extract_docx(self, path: Path) -> Optional[str]:
        """
        Extract text from DOCX file.

        Args:
            path: Path to DOCX file

        Returns:
            str: Extracted text content
        """
        if DocxDocument is None:
            logger.error("python-docx not installed, cannot extract DOCX content")
            return None

        try:
            doc = DocxDocument(str(path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            return '\n\n'.join(paragraphs) if paragraphs else None

        except Exception as e:
            logger.error(f"Error reading DOCX {path.name}: {e}")
            return None

    def _extract_text(self, path: Path) -> Optional[str]:
        """
        Extract content from plain text files.

        Args:
            path: Path to text file

        Returns:
            str: File content
        """
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file {path.name}: {e}")
            return None

    def _extract_unsupported(self, path: Path) -> Optional[str]:
        """
        Placeholder for unsupported formats.

        Args:
            path: Path to file

        Returns:
            None
        """
        logger.warning(f"Format {path.suffix} has limited support, skipping {path.name}")
        return None

    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing excessive whitespace and normalizing.

        Args:
            text: Raw extracted text

        Returns:
            str: Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip()

    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """
        Split text into overlapping chunks for embedding.

        This implements a sliding window chunking strategy optimized for semantic search.
        Overlap ensures context is preserved across chunk boundaries.

        Args:
            text: Text to chunk
            chunk_size: Target size of each chunk (in words)
            overlap: Number of words to overlap between chunks

        Returns:
            List[str]: List of text chunks

        Example:
            >>> extractor = ContentExtractor()
            >>> chunks = extractor.chunk_text("Very long document...", chunk_size=500)
            >>> len(chunks)
            12
        """
        if not text or not text.strip():
            return []

        # Split into words
        words = text.split()

        if len(words) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(words):
            # Extract chunk
            end = start + chunk_size
            chunk_words = words[start:end]
            chunk = ' '.join(chunk_words)
            chunks.append(chunk)

            # Move to next chunk with overlap
            start = end - overlap

            # Prevent infinite loop
            if start >= len(words) - overlap:
                break

        logger.info(f"Split text into {len(chunks)} chunks (size={chunk_size}, overlap={overlap})")
        return chunks


class DocumentIndexer:
    """
    Handles the full pipeline: extraction → chunking → embedding → storage.

    This class orchestrates the RAG content indexing process.
    """

    def __init__(self, vector_db, embedding_service):
        """
        Initialize the document indexer.

        Args:
            vector_db: ChromaDB collection for storing document embeddings
            embedding_service: Service for generating embeddings
        """
        self.extractor = ContentExtractor()
        self.vector_db = vector_db
        self.embedding_service = embedding_service

    def index_document(self, file_path: str, dataset_id: str, doc_id: str) -> int:
        """
        Extract, chunk, embed, and index a document.

        Args:
            file_path: Path to the document
            dataset_id: Parent dataset UUID
            doc_id: Document UUID

        Returns:
            int: Number of chunks indexed

        Example:
            >>> indexer = DocumentIndexer(chroma_collection, embedding_svc)
            >>> count = indexer.index_document("/path/to/doc.pdf", "uuid-123", "doc-456")
            >>> print(f"Indexed {count} chunks")
        """
        # Check if extraction is supported
        if not self.extractor.can_extract(file_path):
            logger.info(f"Skipping {Path(file_path).name} (unsupported format)")
            return 0

        # Extract text content
        content = self.extractor.extract(file_path)

        if not content or len(content) < 100:
            logger.warning(f"Insufficient content extracted from {Path(file_path).name}")
            return 0

        # Chunk the content
        chunks = self.extractor.chunk_text(content, chunk_size=500, overlap=50)

        if not chunks:
            return 0

        # Generate embeddings for each chunk
        embeddings = []
        chunk_texts = []
        metadata_list = []
        ids = []

        for idx, chunk in enumerate(chunks):
            try:
                # Generate embedding
                embedding = self.embedding_service.encode(chunk)

                embeddings.append(embedding.tolist())
                chunk_texts.append(chunk)

                # Metadata for retrieval
                metadata_list.append({
                    'dataset_id': dataset_id,
                    'document_id': doc_id,
                    'chunk_index': idx,
                    'source_file': Path(file_path).name,
                    'type': 'document_content'
                })

                ids.append(f"{doc_id}_chunk_{idx}")

            except Exception as e:
                logger.error(f"Error embedding chunk {idx}: {e}")
                continue

        # Store in ChromaDB
        if embeddings:
            try:
                self.vector_db.add(
                    embeddings=embeddings,
                    documents=chunk_texts,
                    metadatas=metadata_list,
                    ids=ids
                )

                logger.info(f"Indexed {len(embeddings)} chunks from {Path(file_path).name}")
                return len(embeddings)

            except Exception as e:
                logger.error(f"Error storing embeddings in ChromaDB: {e}")
                return 0

        return 0
